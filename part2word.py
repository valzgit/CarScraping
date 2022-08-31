from database import DatabaseInteractor
from docx import Document

document = Document()
document.add_heading('Zadatak 2', 0)
database = DatabaseInteractor()
database.initConnection()

document.add_paragraph("-----------------------Broj po marci-----------------------").add_run("\n" + str(database.printCarNumberPerType()))
document.add_paragraph("-----------------------Broj po lokaciji-----------------------").add_run("\n" + str(database.printCarNumberPerCity()))
document.add_paragraph("-----------------------Broj po boji-----------------------").add_run("\n" + str(database.printCarNumberPerColor()))
document.add_paragraph("-----------------------Rang lista top 30-----------------------").add_run("\n" + str(database.printMostExpensiveCars(30, None)))
document.add_paragraph("-----------------------Rang lista top 30 Dzipova-----------------------").add_run("\n" + str(database.printMostExpensiveCars(30, "Džip/SUV")))
document.add_paragraph("-----------------------Rang lista svih izmedju 2021. i 2022.-----------------------").add_run("\n" + str(database.printRangListBetween20212022()))
document.add_paragraph("-----------------------Najveca kubikaza-----------------------").add_run("\n" + str(database.najvecaKubikaza()))
document.add_paragraph("-----------------------Najveca snaga-----------------------").add_run("\n" + str(database.najvecaSnaga()))
document.add_paragraph("-----------------------Najveca kilometraza-----------------------").add_run("\n" + str(database.najvecaKilometraza()))

document.save('zadatak2.docx')
database.closeConnection()
